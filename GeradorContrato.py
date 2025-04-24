import pandas as pd
from docx import Document
import os
import streamlit as st
from tkinter import Tk, filedialog

def selecionar_arquivo(titulo, tipos):
    root = Tk()
    root.withdraw()
    caminho = filedialog.askopenfilename(title=titulo, filetypes=tipos)
    root.destroy()
    return caminho

def preencher_um_contrato(modelo_path, planilha_path):
    # Lê a planilha Excel
    df = pd.read_excel(planilha_path, sheet_name=0)

    # Transforma em dicionário
    variaveis = pd.Series(df.iloc[:, 1].values, index=df.iloc[:, 0]).to_dict()

    # Carrega o modelo
    doc = Document(modelo_path)

    # Substitui variáveis no texto
    for paragrafo in doc.paragraphs:
        for chave, valor in variaveis.items():
            paragrafo.text = paragrafo.text.replace(f"{{{chave}}}", str(valor))

    # Substitui variáveis em tabelas também
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for chave, valor in variaveis.items():
                    celula.text = celula.text.replace(f"{{{chave}}}", str(valor))

    return doc

# Interface Streamlit
st.title("Preenchimento Automático de Contratos")

st.write(
    "Este aplicativo permite preencher automaticamente um contrato com base em um modelo Word e uma planilha Excel contendo variáveis e valores."
)

modelo_file = st.file_uploader("Escolha o modelo Word com as variáveis", type="docx")
planilha_file = st.file_uploader("Escolha a planilha Excel com as variáveis e valores", type="xlsx")

if modelo_file and planilha_file:
    st.write("Gerando contrato...")
    
    # Salva os arquivos temporários
    modelo_path = "modelo_temp.docx"
    planilha_path = "planilha_temp.xlsx"
    
    with open(modelo_path, "wb") as f:
        f.write(modelo_file.getbuffer())
    
    with open(planilha_path, "wb") as f:
        f.write(planilha_file.getbuffer())
    
    # Preenche o contrato
    doc = preencher_um_contrato(modelo_path, planilha_path)
    
    # Salva o documento gerado
    caminho_saida = "contrato_gerado.docx"
    doc.save(caminho_saida)

    st.write("Contrato gerado com sucesso!")
    st.download_button(
        label="Baixar contrato gerado",
        data=open(caminho_saida, "rb").read(),
        file_name=caminho_saida,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
